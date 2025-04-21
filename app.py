# from flask import Flask, render_template, request, send_file
# from flask_cors import CORS
# from docx import Document
# from docx.shared import Cm, Pt, RGBColor
# import base64
# import os

# app = Flask(__name__)
# CORS(app)



# # Function to add a heading label (for a field name)
# def set_heading_style(paragraph, label):
#     run = paragraph.add_run(f"{label}")
#     run.bold = True
#     run.font.size = Pt(11)
#     run.font.color.rgb = RGBColor(192, 0, 0)  # Dark Red
#     run.font.name = "Calibri"

# # Function to add content text
# def set_content_style(paragraph, text, bold=False):
#     run = paragraph.add_run(text)
#     run.font.name = "Calibri"
#     run.font.size = Pt(11)
#     run.font.color.rgb = RGBColor(0, 0, 0)  # Black
#     run.bold = bold  # Make text bold if needed

# # Function to safely convert string to float
# def safe_float(value, default):
#     try:
#         return float(value) if value and str(value).strip() else default
#     except ValueError:
#         return default

# @app.route('/generate-report', methods=['POST'])
# def generate_report():
#     data = request.json.get('data', [])

#     # Create a new Word document
#     document = Document()
#     document.add_heading('Vulnerability Report', level=1)

#     for vul in data:
#         # --- Title Section ---
#         title = vul.get('Title', 'Unknown')
#         p = document.add_paragraph()
#         set_heading_style(p, "Title")

#         title_para = document.add_paragraph()
#         title_run = title_para.add_run(title)
#         title_run.bold = True
#         title_run.font.size = Pt(13)
#         title_run.font.name = "Devil Breeze bold"
#         title_run.font.color.rgb = RGBColor(89, 89, 89)  # Black

#         # --- Textual Details Section ---
#         fields = [
#             ("Affected Assets", vul.get('Affected_Assets', 'N/A')),
#             ("Description", vul.get('Description', 'N/A')),
#             ("Impact", vul.get('Impact', 'N/A')),
#             ("Recommendations", vul.get('Recommendations', 'N/A')),
#             ("Reference", vul.get('Reference', 'N/A')),
#             ("CVE/CWE", vul.get('CVE_CWE', 'N/A')),
#             ("Status", vul.get('Status', 'N/A')),
#         ]

#         for label, content in fields:
#             p = document.add_paragraph()
#             set_heading_style(p, label)
#             cp = document.add_paragraph()
#             set_content_style(cp, content)

#         # --- Proof of Concept (PoC) Section ---
#         if "PoC" in vul:
#             document.add_page_break()  # Start a new page for PoC
#             p = document.add_paragraph()
#             set_heading_style(p, "Proof of Concept")

#             remaining_space = 25  # Estimated available space per page
#             first_image_on_page = True  # Track first image on a page

#             for i, img_data in enumerate(vul["PoC"]["images"]):
#                 step_text = vul["PoC"]["steps"][i] if i < len(vul["PoC"]["steps"]) else ""

#                 # Ensure Step is included before each image
#                 if step_text:
#                     step_para = document.add_paragraph()
#                     set_content_style(step_para, f"Step {i+1}: ", bold=True)  # Step title in bold
#                     set_content_style(step_para, step_text)  # Step description in normal text

#                 # Decode and save the image
#                 image_path = f"poc_image_{i}.png"
#                 with open(image_path, "wb") as img_file:
#                     img_file.write(base64.b64decode(img_data))

#                 # Get user-defined size, or use default values
#                 width_cm = safe_float(vul["PoC"]["sizes"][i].get("width", ""), 15.9)
#                 height_cm = safe_float(vul["PoC"]["sizes"][i].get("height", ""), 7.73)

#                 # Handle new page option correctly
#                 new_page = vul["PoC"].get("new_page_options", {}).get(str(i), False)  # Default: False

#                 if new_page or remaining_space < height_cm:
#                     # Start a new page if user wants or not enough space
#                     document.add_page_break()
#                     remaining_space = 25  # Reset available space on a new page
#                     first_image_on_page = True  # Reset flag

#                 # Insert the image
#                 document.add_picture(image_path, width=Cm(width_cm), height=Cm(height_cm))

#                 # Reduce available space
#                 remaining_space -= height_cm

#                 # Ensure that the next PoC does not automatically go to a new page
#                 first_image_on_page = False

#                 # Delete the image after adding it to the document
#                 os.remove(image_path)

#         document.add_page_break()  # Ensure each vulnerability starts on a new page

#     # Save the document
#     report_path = "/home/khimeshreport/mysite/vulnerability_report.docx"
#     document.save(report_path)

#     return send_file(report_path, as_attachment=True)

# if __name__ == "__main__":
#     app.run(debug=True)




from flask import Flask, render_template_string, request, send_file
from flask import Flask, request, jsonify, render_template, send_file, redirect, url_for, render_template_string
from flask_cors import CORS
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import base64
import os

app = Flask(__name__)
CORS(app)

# Style: Heading label
def set_heading_style(paragraph, label):
    run = paragraph.add_run(label)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(192, 0, 0)  # Dark Red
    run.font.name = "Calibri"

# Style: Content text
def set_content_style(paragraph, text, bold=False):
    run = paragraph.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.bold = bold

# Safe float conversion
def safe_float(value, default):
    try:
        return float(value) if value and str(value).strip() else default
    except ValueError:
        return default

# Add image in a table with black border
def add_styled_image(doc, image_path, width_cm, height_cm):
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)

    tbl = table._element
    tblBorders = OxmlElement('w:tblBorders')

    for border_name in ["top", "left", "bottom", "right"]:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:color"), "000000")
        tblBorders.append(border)

    tbl.append(tblBorders)

    run = cell.paragraphs[0].add_run()
    run.add_picture(image_path, width=Cm(width_cm), height=Cm(height_cm))

    doc.add_paragraph("\n")  # Spacing below image

# Create clickable hyperlink in Word
def create_hyperlink(paragraph, url, text):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True
    )

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # Blue color and underline
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0000FF')
    rPr.append(color)

    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)

    new_run.append(rPr)

    text_elem = OxmlElement('w:t')
    text_elem.text = text
    new_run.append(text_elem)
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

# preview area

@app.route('/preview-report', methods=['POST'])
def preview_report():
    data = request.json.get('data', [])
    html = "<h1>Vulnerability Report Preview</h1>"
    for idx, vul in enumerate(data, 1):
        html += f"<h2>{idx}. {vul.get('Title', 'No Title')}</h2>"
        html += f"<p><strong>Description:</strong> {vul.get('Description', '')}</p>"
        html += f"<p><strong>Impact:</strong> {vul.get('Impact', '')}</p>"
        html += f"<p><strong>Recommendations:</strong> {vul.get('Recommendations', '')}</p>"
        html += f"<p><strong>Reference:</strong> {vul.get('Reference', '')}</p>"
        html += f"<p><strong>CVE/CWE:</strong> {vul.get('CVE_CWE', '')}</p>"
        html += f"<p><strong>Status:</strong> {vul.get('Status', '')}</p>"

        
        # Display images with steps
        poc = vul.get('PoC', {})
        images = poc.get('images', [])
        steps = poc.get('steps', [])
        sizes = poc.get('sizes', [])
        if images:
            html += "<h3>Proof of Concept:</h3>"
            for i, img_data in enumerate(images):
                step = steps[i] if i < len(steps) else ""
                size = sizes[i] if i < len(sizes) else {}
                width = size.get('width', 'auto')
                height = size.get('height', 'auto')
                html += f"<p><strong>Step {i+1}:</strong> {step}</p>"
                html += f'<img src="data:image/png;base64,{img_data}" style="max-width:100%;width:{width}cm;height:{height}cm;border:1px solid #000;margin-bottom:10px;" /><br>'
    return render_template_string(html)

@app.route('/generate-report', methods=['POST'])
def generate_report():
    data = request.json.get('data', [])
    document = Document()
    document.add_heading('Vulnerability Report', level=1)

    for vul in data:
        # --- Title ---
        title = vul.get('Title', 'Unknown')
        p = document.add_paragraph()
        set_heading_style(p, "Title")
        p_title = document.add_paragraph()
        title_run = p_title.add_run(title)
        title_run.bold = True
        title_run.font.size = Pt(13)
        title_run.font.name = "Devil Breeze"
        title_run.font.color.rgb = RGBColor(89, 89, 89)

        # --- Fields ---
        fields = [
            ("Affected Assets", vul.get('Affected_Assets', 'N/A')),
            ("Description", vul.get('Description', 'N/A')),
            ("Impact", vul.get('Impact', 'N/A')),
            ("Recommendations", vul.get('Recommendations', 'N/A')),
            ("Reference", vul.get('Reference', 'N/A')),
            ("CVE/CWE", vul.get('CVE_CWE', 'N/A')),
            ("Status", vul.get('Status', 'N/A')),
        ]

        for label, content in fields:
            p = document.add_paragraph()
            set_heading_style(p, label)
            cp = document.add_paragraph()
            
            # Add hyperlink if content is a URL
            if isinstance(content, str) and content.startswith(('http://', 'https://')):
                create_hyperlink(cp, content, content)
            else:
                set_content_style(cp, content)

        # --- PoC Section ---
        if "PoC" in vul and vul["PoC"]["images"]:
            document.add_page_break()

            poc_heading = document.add_paragraph()
            set_heading_style(poc_heading, "Proof of Concept")

            for i, img_data in enumerate(vul["PoC"]["images"]):
                step_text = vul["PoC"]["steps"][i] if i < len(vul["PoC"]["steps"]) else ""
                if step_text:
                    step_para = document.add_paragraph()
                    set_content_style(step_para, f"Step {i+1}: ", bold=True)
                    set_content_style(step_para, step_text)

                image_path = f"poc_image_{i}.png"
                with open(image_path, "wb") as img_file:
                    img_file.write(base64.b64decode(img_data))

                width_cm = safe_float(vul["PoC"]["sizes"][i].get("width", ""), 15.9)
                height_cm = safe_float(vul["PoC"]["sizes"][i].get("height", ""), 7.73)

                add_styled_image(document, image_path, width_cm, height_cm)
                os.remove(image_path)

        document.add_page_break()

    report_path = "/home/khimeshreport/mysite/vulnerability_report.docx"
    document.save(report_path)

    return send_file(report_path, as_attachment=True)

if __name__ == "__main__":
    app.run(debug=True)




